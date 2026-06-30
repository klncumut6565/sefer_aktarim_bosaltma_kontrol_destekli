#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Boşaltma Kontrol Dökümanı (.docx) otomatik doldurma modülü.

Şablon dosyasındaki (BOŞALTMA_KONTROL_SABLONU.docx) boş alanları,
PDF'ten çıkarılan sefer verileri ve kullanıcının girdiği geçerlilik/muayene
tarihleri ile doldurarak yeni bir .docx üretir.

Kullanım:
    from docx_doldur import kontrol_dokumani_olustur

    kontrol_dokumani_olustur(
        sablon_path="BOŞALTMA_KONTROL_SABLONU.docx",
        cikti_path="Kontrol_12345.docx",
        tarih="01.07.2026",
        gonderici_unvan="ABC Kimya A.Ş.",
        tasiyici_unvan="XYZ Lojistik Ltd.",
        plaka="34 ABC 123",
        sefer_un_listesi="12345 / UN 1203, UN 1202",
        yangin_tup_tarihi="15.08.2026",
        tmfb_tarihi="20.09.2026",
        adr_uygunluk_tarihi="10.10.2026",
        ara_muayene_tarihi="01.01.2027",
        periyodik_muayene_tarihi="01.06.2027",
    )
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

import docx
from docx.table import _Cell


# ---------------------------------------------------------------------------
# Üst tablo (Table 0): Evet / Hayır / İlgili Değil seçim mantığı
# ---------------------------------------------------------------------------
# Madde 1 -> "Evet" işaretlenecek
# Madde 2 -> "İlgili Değil" işaretlenecek
# Madde 3 -> "İlgili Değil" işaretlenecek
UST_SECIMLER = {
    0: "Evet",
    1: "İlgili Değil",
    2: "İlgili Değil",
}

# Sol tablo (Alıcı-Dolduran-Boşaltan) - satır indeksleri (Table 2 satır no -> 1'den başlar)
# 2 ve 7 numaralı maddeler zaten "İ.D." / "İ.D" sabit yazılı, dokunulmuyor.
# Diğerleri (1,3,4,5,6,8,9) "✓" ile doldurulacak.
SOL_TIK_SATIRLAR = {2, 4, 5, 6, 7, 9, 10}  # Table 2 içindeki gerçek satır indeksleri (0-bazlı)
SAG_TIK_SATIRLAR = {2, 3, 4, 5, 6, 7, 8, 9, 10}  # sağ taraftaki tüm 1-9 satırları

TIK_ISARETI = "✓"


def _set_cell_text_keep_format(cell: _Cell, text: str, paragraph_index: int = 0) -> None:
    """Hücredeki belirli bir paragrafın metnini, ilk run'ın biçimini koruyarak değiştirir.
    Eğer hücre boşsa, mevcut tek paragrafa yazar."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = text
        return
    p = paragraphs[paragraph_index] if paragraph_index < len(paragraphs) else paragraphs[0]
    if p.runs:
        # İlk run'ın formatını koru, diğer run'ları temizle
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _append_to_cell(cell: _Cell, text: str) -> None:
    """Hücrenin sonuna yeni bir paragraf olarak metin ekler (mevcut içeriği bozmadan)."""
    p = cell.add_paragraph()
    p.add_run(text)


def _replace_tarih_placeholder(cell: _Cell, deger: str) -> bool:
    """Hücre içindeki 'Geçerlilik Tarihi:' / 'Ara Muayene tarihi:' / 'Periyodik Muayene
    tarihi:' satırındaki nokta/eğik çizgi placeholder'ını verilen tarihle değiştirir.
    Run'lar kelime ortasında bölünmüş olabileceğinden paragraf seviyesinde işlem yapılır.
    Tek hücrede birden fazla satır (örn. Ara + Periyodik Muayene) olabilir; tüm eşleşen
    paragraflar bulunur ve sırasıyla doldurulur. `deger` bir liste de olabilir.
    """
    degerler = deger if isinstance(deger, list) else [deger]
    idx = 0
    degisti = False
    for p in cell.paragraphs:
        full_text = "".join(r.text for r in p.runs) if p.runs else p.text
        if ":" in full_text and ("/" in full_text or "…" in full_text or "." in full_text.split(":")[-1]):
            etiket = full_text.split(":")[0] + ": "
            if idx < len(degerler):
                yeni_metin = etiket + degerler[idx]
                if p.runs:
                    p.runs[0].text = yeni_metin
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    for r in list(p.runs):
                        r.text = ""
                    p.add_run(yeni_metin)
                idx += 1
                degisti = True
    return degisti


def _append_isim_to_cell(cell: _Cell, isim: str) -> None:
    """Footer hücresindeki 'Boşaltan\t:' veya 'Taşıyıcı/Şoför Adı Soyadı:' gibi
    etiketin sonuna, formatı koruyarak ismi ekler (etiket: İSİM SOYİSİM)."""
    if not cell.paragraphs:
        return
    p = cell.paragraphs[0]
    if p.runs:
        son_run = p.runs[-1]
        son_run.text = f"{son_run.text} {isim}"
    else:
        p.add_run(f" {isim}")


def kontrol_dokumani_olustur(
    sablon_path: str | Path,
    cikti_path: str | Path,
    tarih: str = "",
    gonderici_unvan: str = "",
    tasiyici_unvan: str = "",
    plaka: str = "",
    sefer_un_listesi: str = "",
    yangin_tup_tarihi: str = "",
    tmfb_tarihi: str = "",
    adr_uygunluk_tarihi: str = "",
    ara_muayene_tarihi: str = "",
    periyodik_muayene_tarihi: str = "",
    bosaltan_adi: str = "",
    sofor_adi: str = "",
) -> Path:
    """Şablon .docx dosyasını doldurup yeni bir dosya olarak kaydeder.

    Tarih parametreleri boş bırakılırsa ilgili placeholder olduğu gibi kalır
    (kullanıcı Word'de elle doldurabilir).
    """
    sablon_path = Path(sablon_path)
    cikti_path = Path(cikti_path)
    if not sablon_path.is_file():
        raise FileNotFoundError(f"Şablon bulunamadı: {sablon_path}")

    d = docx.Document(str(sablon_path))

    if len(d.tables) < 3:
        raise ValueError(
            f"Beklenmeyen şablon yapısı: {len(d.tables)} tablo bulundu, en az 3 bekleniyordu."
        )

    # ---- Tablo 0: Üst Evet/Hayır/İlgili Değil seçimleri ----
    t0 = d.tables[0]
    for row_idx, secim in UST_SECIMLER.items():
        if row_idx >= len(t0.rows):
            continue
        cell = t0.rows[row_idx].cells[0]
        for p in cell.paragraphs:
            full_text = "".join(r.text for r in p.runs) if p.runs else p.text
            if full_text.strip().startswith("[ ]") and secim in full_text:
                yeni = full_text.replace("[ ]", "[X]", 1) if full_text.strip().startswith(f"[ ] {secim}") else full_text
                # Daha güvenli: sadece ilgili seçimin başındaki [ ] işaretini değiştir
        # Basit ve güvenilir yöntem: tüm paragraf metnini yeniden kur
        for p in cell.paragraphs:
            full_text = "".join(r.text for r in p.runs) if p.runs else p.text
            if "[ ]" not in full_text:
                continue
            yeni_metin = full_text.replace(f"[ ] {secim}", f"[X] {secim}")
            if yeni_metin != full_text:
                if p.runs:
                    p.runs[0].text = yeni_metin
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.add_run(yeni_metin)

    # ---- Tablo 1: Tarih / Firma / Plaka / Sefer bilgileri (5x2) ----
    t1 = d.tables[1]
    deger_map = {
        0: tarih,
        1: gonderici_unvan,
        2: tasiyici_unvan,
        3: plaka,
        4: sefer_un_listesi,
    }
    for row_idx, deger in deger_map.items():
        if not deger or row_idx >= len(t1.rows):
            continue
        hedef_cell = t1.rows[row_idx].cells[1]
        _set_cell_text_keep_format(hedef_cell, str(deger))

    # ---- Tablo 2: Sol/Sağ kontrol listeleri (11x6) ----
    t2 = d.tables[2]

    for row_idx in SOL_TIK_SATIRLAR:
        if row_idx >= len(t2.rows):
            continue
        onay_cell = t2.rows[row_idx].cells[2]
        _set_cell_text_keep_format(onay_cell, TIK_ISARETI)

    for row_idx in SAG_TIK_SATIRLAR:
        if row_idx >= len(t2.rows):
            continue
        onay_cell = t2.rows[row_idx].cells[5]
        _set_cell_text_keep_format(onay_cell, TIK_ISARETI)

    # ---- Tarih placeholder'ları (sol tablo, sütun index 1) ----
    # Satır 2 (görsel "1"): Araçta Yangın Tüpü -> sağ tabloda! sütun index 4
    if yangin_tup_tarihi:
        _replace_tarih_placeholder(t2.rows[2].cells[4], yangin_tup_tarihi)
    # Satır 5 (görsel "4"): Taşımacı TMFB -> sol tablo, sütun index 1
    if tmfb_tarihi:
        _replace_tarih_placeholder(t2.rows[5].cells[1], tmfb_tarihi)
    # Satır 6 (görsel "5"): ADR Uygunluk Belgesi -> sol tablo, sütun index 1
    if adr_uygunluk_tarihi:
        _replace_tarih_placeholder(t2.rows[6].cells[1], adr_uygunluk_tarihi)
    # Satır 7 (görsel "6"): Ara Muayene + Periyodik Muayene -> sol tablo, sütun index 1 (2 satır)
    if ara_muayene_tarihi or periyodik_muayene_tarihi:
        degerler = []
        if ara_muayene_tarihi:
            degerler.append(ara_muayene_tarihi)
        if periyodik_muayene_tarihi:
            # İkisi de varsa sırayla, sadece biri varsa onu tek başına eşle
            if ara_muayene_tarihi:
                degerler.append(periyodik_muayene_tarihi)
            else:
                degerler = [periyodik_muayene_tarihi]
        _replace_tarih_placeholder(t2.rows[7].cells[1], degerler)

    # ---- Footer: Boşaltan / Taşıyıcı-Şoför Adı Soyadı ----
    if bosaltan_adi or sofor_adi:
        try:
            footer = d.sections[0].footer
            if footer.tables:
                ft = footer.tables[0]
                if bosaltan_adi:
                    _append_isim_to_cell(ft.rows[0].cells[0], bosaltan_adi)
                if sofor_adi:
                    _append_isim_to_cell(ft.rows[0].cells[1], sofor_adi)
        except (IndexError, AttributeError):
            pass

    cikti_path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(cikti_path))
    return cikti_path


if __name__ == "__main__":
    # Hızlı manuel test
    out = kontrol_dokumani_olustur(
        sablon_path="/home/claude/proj/Bosaltma_Kontrol_Sablonu.docx",
        cikti_path="/home/claude/proj/test_cikti.docx",
        tarih="01.07.2026",
        gonderici_unvan="ABC Kimya Sanayi A.Ş.",
        tasiyici_unvan="XYZ Lojistik Taşımacılık Ltd. Şti.",
        plaka="34 ABC 123",
        sefer_un_listesi="Sefer No: 12345 — UN 1203, UN 1202",
        yangin_tup_tarihi="15.08.2026",
        tmfb_tarihi="20.09.2026",
        adr_uygunluk_tarihi="10.10.2026",
        ara_muayene_tarihi="01.01.2027",
        periyodik_muayene_tarihi="01.06.2027",
        bosaltan_adi="Mehmet Yılmaz",
        sofor_adi="Ahmet Demir",
    )
    print("Oluşturuldu:", out)
