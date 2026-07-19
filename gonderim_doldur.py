#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gönderim Kontrol Dökümanı (.docx) otomatik doldurma modülü.

Boşaltma modülündeki docx_doldur.py ile aynı yaklaşımı kullanır:
- Şablon DOCX açılır, boş hücreler/checkbox'lar doldurulur, PDF'e çevrilir.
- Tüm checkbox'lar ✓ ile doldurulur.
- Tarih, Gönderici Firma, Taşıyıcı, Plaka, Taşıma Evrağı No export verilerinden gelir.
- Araç Türü: her zaman Kamyon (ambalajlı atık taşıması).
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

import docx

TIK = '✓'

# Üst tablo checkbox seçimleri — atık gönderimleri her zaman ADR-AMBALAJLI
# Soru 1 (Ambalajlı plaka): Evet   |   Soru 2 (Tank/Dökme plaka): İlgili Değil
# Soru 3 (Konteynır): İlgili Değil
UST_SECIMLER = {0: 'Evet', 1: 'İlgili Değil', 2: 'İlgili Değil'}

# Gönderen-Paketleyen kontrolleri: tüm satırlar ✓
GONDEREN_TIK_SATIRLAR = {2, 3, 4, 5, 6, 7, 8, 9, 10}   # Table 2, sol taraf (satır indeksi)
# Yükleyen kontrolleri: tüm satırlar ✓
YUKLEYEN_TIK_SATIRLAR = {2, 3, 4, 5, 6, 7, 8, 9, 10}   # Table 2, sağ taraf

# Araç Türü Kamyon checkbox metni (şablondaki tam text'e göre)
ARAC_TURU_KAMYON = '[ ] Kamyon'


def _write_cell(cell, text: str) -> None:
    """Hücrenin ilk paragrafını formatı koruyarak yazar."""
    p = cell.paragraphs[0] if cell.paragraphs else None
    if p is None:
        return
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(text)


def _fill_checkbox(cell, secim: str) -> None:
    """
    Hücre içindeki [ ] Evet / [ ] Hayır / [ ] İlgili Değil satırlarından
    secim ile eşleşeni [X] olarak işaretler.
    """
    for p in cell.paragraphs:
        full = ''.join(r.text for r in p.runs) if p.runs else p.text
        if '[ ]' not in full:
            continue
        if secim in full:
            yeni = full.replace(f'[ ] {secim}', f'[X] {secim}', 1)
            if p.runs:
                p.runs[0].text = yeni
                for r in p.runs[1:]:
                    r.text = ''
            else:
                p.clear()
                p.add_run(yeni)


def _logo_degistir(d: docx.Document, logo_bytes: bytes) -> None:
    """Belgedeki ilk logo resmini (rId7/image1.png — sol üst köşe) kullanıcı logosuyla değiştirir.
    Gönderim şablonunda logo body'deki ilk tablonun hücresinde, header'da değil."""
    IMAGE_REL_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'

    # Önce document part rels içinde dene (body'deki resimler buradan)
    for rel in d.part.rels.values():
        if rel.reltype == IMAGE_REL_TYPE:
            # En büyük resmi logo kabul et (ikonlar küçük, logo büyük)
            try:
                mevcut = rel.target_part._blob
                from PIL import Image
                import io as _io
                im = Image.open(_io.BytesIO(mevcut))
                if im.size[0] > 200:  # 200px'den geniş → logo
                    rel.target_part._blob = logo_bytes
                    return
            except Exception:
                rel.target_part._blob = logo_bytes
                return

    # Fallback: header rels
    for section in d.sections:
        try:
            for rel in section.header.part.rels.values():
                if rel.reltype == IMAGE_REL_TYPE:
                    rel.target_part._blob = logo_bytes
                    return
        except AttributeError:
            pass


def gonderim_dokumani_olustur(
    sablon_path: str | Path,
    cikti_path: str | Path,
    tarih: str = '',
    gonderici_firma: str = '',
    tasiyici_firma: str = '',
    plaka: str = '',
    tasima_evragi_no: str = '',
    gonderici_adi: str = '',
    sofor_adi: str = '',
    logo_bytes: Optional[bytes] = None,
) -> Path:
    """
    Gönderim Kontrol Dökümanı şablonunu doldurup .docx olarak kaydeder.

    Args:
        sablon_path: Şablon .docx dosyasının yolu
        cikti_path: Çıktı .docx dosyasının yolu
        tarih: GG.AA.YYYY formatında tarih
        gonderici_firma: Gönderici/üretici firma unvanı (kullanıcı girer)
        tasiyici_firma: Taşıyıcı firma unvanı (export'tan gelir)
        plaka: Araç plaka numarası
        tasima_evragi_no: Taşıma numarası/numaraları (ör. E8433002, E8432994)
        gonderici_adi: Belgenin altındaki Gönderen adı soyadı
        sofor_adi: Taşıyıcı/Şoför adı soyadı
        logo_bytes: Firma logosu ham bytes (None ise değiştirilmez)
    """
    d = docx.Document(str(sablon_path))

    if len(d.tables) < 3:
        raise ValueError(f'Beklenmeyen şablon yapısı: {len(d.tables)} tablo var')

    # ---- Logo ----
    if logo_bytes:
        _logo_degistir(d, logo_bytes)

    # ---- Table 0: Üst 3 checkbox (Evet/İlgili Değil/İlgili Değil) ----
    t0 = d.tables[0]
    for row_idx, secim in UST_SECIMLER.items():
        if row_idx < len(t0.rows):
            _fill_checkbox(t0.rows[row_idx].cells[0], secim)

    # ---- Table 1: Tarih/Firma/Plaka bilgileri ----
    t1 = d.tables[1]
    bilgi_map = {
        0: tarih,           # Tarih
        1: gonderici_firma, # Gönderici Firma Unvanı
        2: tasiyici_firma,  # Taşıyıcı Firma Unvanı
        3: plaka,           # Araç Plaka No
    }
    for row_idx, deger in bilgi_map.items():
        if deger and row_idx < len(t1.rows):
            _write_cell(t1.rows[row_idx].cells[1], deger)

    # Satır 4: Araç Türü — Kamyon seçili
    if len(t1.rows) > 4:
        arac_cell = t1.rows[4].cells[1]
        mevcut = ''.join(r.text for p in arac_cell.paragraphs for r in p.runs)
        if '[ ] Kamyon' in mevcut:
            yeni = mevcut.replace('[ ] Kamyon', '[X] Kamyon', 1)
            for p in arac_cell.paragraphs:
                for r in p.runs:
                    r.text = ''
            arac_cell.paragraphs[0].add_run(yeni)

    # ---- Table 2: Gönderen-Paketleyen / Yükleyen onay kutucukları ----
    t2 = d.tables[2]
    for row_idx in GONDEREN_TIK_SATIRLAR:
        if row_idx < len(t2.rows):
            _write_cell(t2.rows[row_idx].cells[2], TIK)
    for row_idx in YUKLEYEN_TIK_SATIRLAR:
        if row_idx < len(t2.rows):
            _write_cell(t2.rows[row_idx].cells[5], TIK)

    # ---- Table 4: İmza tablosu (body'de, footer değil) ----
    # Gönderim formunda imza bilgileri body'deki son tabloda
    for ti in range(len(d.tables) - 1, -1, -1):
        t = d.tables[ti]
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if 'Gönderen' in txt and ri == 0 and ci == 0:
                    # İmza satırı bulundu
                    if gonderici_adi:
                        # "Gönderen" yazısının yanına isim ekle
                        for p in cell.paragraphs:
                            full = ''.join(r.text for r in p.runs)
                            if 'Gönderen' in full and ':' in full:
                                if p.runs:
                                    p.runs[-1].text = p.runs[-1].text + f' {gonderici_adi}'
                                break
                    if sofor_adi and ci + 1 < len(row.cells):
                        sof_cell = row.cells[ci + 1]
                        for p in sof_cell.paragraphs:
                            full = ''.join(r.text for r in p.runs)
                            if 'Şoför' in full or 'Soyadı' in full:
                                if p.runs:
                                    p.runs[-1].text += f' {sofor_adi}'
                                break

    # Taşıma Evrağı No'yu talimat tablosunun üstüne ekle (paragraf olarak)
    # veya Table 1'e ekstra satır olarak — burada Table 1'deki 5. satırı kullan
    # Şablonda Table 1 sadece 5 satır (0-4), Araç Türü 4. satır
    # Taşıma No'yu Araç Türü altında yeni satır olarak değil,
    # header/title alanına yazıyoruz (belge başlığı altında zaten var)
    # Bunun yerine: Table 1 satır 4 (Araç Türü) sonrasına programatik eklemek
    # karmaşık, bu yüzden talimat metnine (Table 3) ilk satır olarak ekliyoruz.
    # Daha temiz çözüm: Taşıma No'yu "Taşıma Evrağı Numarası" olarak
    # zaten var olan Table 1 Satır 4'e (değer hücresine) yazıyoruz
    # Ama Araç Türü checkbox'larını bozmamak için ayrı paragraf ekle
    if tasima_evragi_no and len(t1.rows) > 4:
        # Taşıma Evrağı No'yu talimat tabloya ilk paragraf olarak ekle
        if len(d.tables) >= 4:
            t3 = d.tables[3]
            if t3.rows:
                ilk_cell = t3.rows[0].cells[0]
                # Mevcut içeriğin önüne ekle
                p_yeni = ilk_cell.paragraphs[0]._element
                from docx.oxml import OxmlElement
                p_no = OxmlElement('w:p')
                r_no = OxmlElement('w:r')
                t_no = OxmlElement('w:t')
                t_no.text = f'Taşıma Evrağı No: {tasima_evragi_no}'
                r_no.append(t_no)
                p_no.append(r_no)
                ilk_cell._element.insert(0, p_no)

    cikti_path = Path(cikti_path)
    cikti_path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(cikti_path))
    return cikti_path
